import requests
from lxml import etree
from pymongo import MongoClient
import docx


class Poetry:
    def __init__(self, t, h, c, s, l):
        self.title = t
        self.auth = h
        self.content = c
        self.dynasty = s
        self.label = l

    def __str__(self):
        content = """
        题目：{}
        作者：{}
        朝代：{}
        内容：
            {}
        标签：{}
        """.format(self.title, self.auth, self.dynasty, self.content, self.label)
        return content

    def save_mongo(self):
        conn = MongoClient("localhost", 27017)
        db = conn.SpiderData
        poetry = db.Poetry
        poetry.save({
            "title": self.title,
            "auth": self.auth,
            "dynasty": self.dynasty,
            "content": self.content,
            "label": self.label
        })


URL = ["https://www.gushiwen.org/shiwen/default.aspx?page=%d&type=0&id=0" % p for p in range(1, 1001)]


def get_content(url):
    html = requests.get(url)
    html = etree.HTML(html.text)
    lis = html.xpath('//div[@class="left"]/div[@class="sons"]')
    for l in lis:
        title = l.xpath("./div/p//a/b/text()")[0]
        auth = l.xpath("./div/p[2]/a[2]/text()")[0]
        dynasty = l.xpath("./div/p[2]/a[1]/text()")[0]

        contentList = l.xpath("./div/div[2]//text()")
        labelList = l.xpath("./div[@class='tag']//text()")
        labelList = list(set(labelList))
        try:
            labelList.remove("\n")
            labelList.remove("，")
        except Exception as e:
            print(e)
        label = tuple(set(labelList))
        content = "\n".join(contentList)
        p = Poetry(title, auth, content, dynasty, label)
        p.save_mongo()
        print(p.title)


"""
1 生成一级目录：朝代
2.生成二级目录：作者
3.生成三级目录：诗题
4.写入内容：诗内容
"""


def export_word():
    doc = docx.Document()
    conn = MongoClient("localhost", 27017)
    db = conn.SpiderData
    poetry = db.Poetry
    dynastyLis = ['先秦', '两汉', '魏晋', '南北朝', '唐代', '五代', '金朝', '宋代', '明代', '元代', '清代', '近现代', '未知']
    for dynasty in dynastyLis:
        doc.add_heading(dynasty, level=1)

        # 获取诗人集合
        authLis = set()
        for a in poetry.find({"dynasty": dynasty}):
            authLis.add(a["auth"])

        for auth in authLis:
            doc.add_heading(auth, level=2)

            # 获取诗人的所有诗
            titleList = set()
            for title in poetry.find({"dynasty": dynasty, "auth": auth}):
                titleList.add(title["title"])

            for t in titleList:
                doc.add_heading(t, level=3)

                result = poetry.find_one({"dynasty": dynasty, "auth": auth, "title": t})
                c = doc.add_paragraph(result["content"])
    doc.save("poetry.docx")


def run():
    for u in URL:
        get_content(u)
        # time.sleep(5)


if __name__ == '__main__':
    export_word()
